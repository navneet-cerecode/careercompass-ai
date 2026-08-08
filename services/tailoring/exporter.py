"""Deterministic DOCX and PDF exports for user-verified resume content."""

from __future__ import annotations

from html import escape
from io import BytesIO
from typing import Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

from models.tailored_resume import TailoredResumeContent

ExportFormat = Literal["docx", "pdf"]

INK = RGBColor(16, 35, 31)
MUTED = RGBColor(77, 94, 89)
ACCENT = RGBColor(57, 126, 86)


class ResumeDocumentExporter:
    """Render a compact resume without changing any accepted content."""

    def render(self, content: TailoredResumeContent, export_format: ExportFormat) -> bytes:
        if export_format == "docx":
            return self._render_docx(content)
        if export_format == "pdf":
            return self._render_pdf(content)
        raise ValueError(f"Unsupported export format: {export_format}")

    def _render_docx(self, content: TailoredResumeContent) -> bytes:
        document = Document()
        section = document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.65)
        section.right_margin = Inches(0.72)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.72)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)

        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = "Arial"
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        normal.font.size = Pt(9.5)
        normal.font.color.rgb = INK
        normal.paragraph_format.space_after = Pt(3)
        normal.paragraph_format.line_spacing = 1.08

        bullet = styles["List Bullet"]
        bullet.font.name = "Arial"
        bullet._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        bullet.font.size = Pt(9.5)
        bullet.paragraph_format.left_indent = Inches(0.2)
        bullet.paragraph_format.first_line_indent = Inches(-0.14)
        bullet.paragraph_format.space_after = Pt(2.5)
        bullet.paragraph_format.line_spacing = 1.06

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(2)
        run = title.add_run(content.name)
        self._set_run_font(run, size=20, color=INK, bold=True)

        contact_lines = self._contact_lines(content)
        for index, line in enumerate(contact_lines):
            contact = document.add_paragraph()
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact.paragraph_format.space_after = Pt(7 if index == len(contact_lines) - 1 else 1)
            contact_run = contact.add_run(line)
            self._set_run_font(contact_run, size=8.5, color=MUTED)

        if content.skills:
            self._add_section_heading(document, "Skills")
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.add_run(", ".join(skill.name for skill in content.skills))

        self._add_list_section(document, "Experience", content.experience)
        self._add_list_section(document, "Projects", content.projects)
        self._add_list_section(document, "Education", content.education)
        self._add_list_section(document, "Certifications", content.certifications)
        self._add_list_section(document, "Achievements", content.achievements)

        document.core_properties.author = ""
        document.core_properties.last_modified_by = ""
        document.core_properties.title = "Resume"
        stream = BytesIO()
        document.save(stream)
        return stream.getvalue()

    @staticmethod
    def _set_run_font(
        run,
        *,
        size: float,
        color: RGBColor,
        bold: bool = False,
    ) -> None:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold

    def _add_section_heading(self, document: Document, title: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(title.upper())
        self._set_run_font(run, size=10.5, color=ACCENT, bold=True)
        p_pr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "B9EAD5")
        borders.append(bottom)
        p_pr.append(borders)

    def _add_list_section(
        self,
        document: Document,
        title: str,
        items: tuple[str, ...],
    ) -> None:
        if not items:
            return
        self._add_section_heading(document, title)
        for item in items:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(item)

    @staticmethod
    def _contact_lines(content: TailoredResumeContent) -> tuple[str, ...]:
        primary = " | ".join(value for value in (content.email, content.phone) if value)
        links = " | ".join(value for value in (content.linkedin, content.github) if value)
        return tuple(line for line in (primary, links) if line)

    @staticmethod
    def _render_pdf(content: TailoredResumeContent) -> bytes:
        stream = BytesIO()
        document = SimpleDocTemplate(
            stream,
            pagesize=letter,
            rightMargin=0.72 * inch,
            leftMargin=0.72 * inch,
            topMargin=0.65 * inch,
            bottomMargin=0.65 * inch,
            title="Resume",
            author="",
        )
        styles = getSampleStyleSheet()
        body = ParagraphStyle(
            "ResumeBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=11.2,
            textColor=colors.HexColor("#10231F"),
            spaceAfter=3,
            alignment=TA_LEFT,
        )
        title = ParagraphStyle(
            "ResumeTitle",
            parent=body,
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=22,
            spaceAfter=2,
            alignment=TA_CENTER,
        )
        contact = ParagraphStyle(
            "ResumeContact",
            parent=body,
            fontSize=8.5,
            leading=10,
            textColor=colors.HexColor("#4D5E59"),
            alignment=TA_CENTER,
            spaceAfter=1,
        )
        heading = ParagraphStyle(
            "ResumeHeading",
            parent=body,
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=12,
            textColor=colors.HexColor("#397E56"),
            spaceBefore=6,
            spaceAfter=3,
            borderWidth=0,
            borderPadding=0,
        )

        story = [Paragraph(escape(content.name), title)]
        contact_lines = ResumeDocumentExporter._contact_lines(content)
        for line in contact_lines:
            story.append(Paragraph(escape(line), contact))
        if contact_lines:
            story.append(Spacer(1, 4))

        if content.skills:
            story.extend(
                [
                    Paragraph("SKILLS", heading),
                    Paragraph(
                        escape(", ".join(skill.name for skill in content.skills)),
                        body,
                    ),
                ]
            )

        for section_title, items in (
            ("EXPERIENCE", content.experience),
            ("PROJECTS", content.projects),
            ("EDUCATION", content.education),
            ("CERTIFICATIONS", content.certifications),
            ("ACHIEVEMENTS", content.achievements),
        ):
            if not items:
                continue
            story.append(Paragraph(section_title, heading))
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(escape(item), body), leftIndent=9) for item in items],
                    bulletType="bullet",
                    start="circle",
                    leftIndent=14,
                    bulletFontName="Helvetica",
                    bulletFontSize=6,
                    spaceAfter=2,
                )
            )
            story.append(Spacer(1, 2))

        document.build(story)
        return stream.getvalue()
