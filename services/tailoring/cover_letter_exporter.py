"""DOCX and PDF exports for user-verified cover letters."""

from html import escape
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from models.cover_letter import CoverLetterContent
from services.tailoring.exporter import ExportFormat

INK = RGBColor(16, 35, 31)
MUTED = RGBColor(77, 94, 89)


class CoverLetterDocumentExporter:
    """Render accepted cover letter text without rewriting it."""

    def render(self, content: CoverLetterContent, export_format: ExportFormat) -> bytes:
        if export_format == "docx":
            return self._render_docx(content)
        if export_format == "pdf":
            return self._render_pdf(content)
        raise ValueError(f"Unsupported export format: {export_format}")

    @staticmethod
    def _render_docx(content: CoverLetterContent) -> bytes:
        document = Document()
        section = document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.85)
        section.right_margin = Inches(0.9)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.9)

        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        normal.font.size = Pt(11)
        normal.font.color.rgb = INK
        normal.paragraph_format.space_after = Pt(10)
        normal.paragraph_format.line_spacing = 1.15

        name = document.add_paragraph()
        name.paragraph_format.space_after = Pt(2)
        run = name.add_run(content.candidate_name)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = INK
        if content.candidate_email:
            contact = document.add_paragraph(content.candidate_email)
            contact.paragraph_format.space_after = Pt(14)
            contact.runs[0].font.size = Pt(9)
            contact.runs[0].font.color.rgb = MUTED

        role = document.add_paragraph()
        role.paragraph_format.space_after = Pt(14)
        role_run = role.add_run(
            f"{content.job_title} at {content.company_name}"
        )
        role_run.bold = True

        for paragraph_text in (
            content.salutation,
            content.opening,
            content.evidence_paragraph,
            content.motivation_paragraph,
            content.closing_paragraph,
            content.sign_off,
            content.candidate_name,
        ):
            document.add_paragraph(paragraph_text)

        document.core_properties.author = ""
        document.core_properties.last_modified_by = ""
        document.core_properties.title = "Cover Letter"
        stream = BytesIO()
        document.save(stream)
        return stream.getvalue()

    @staticmethod
    def _render_pdf(content: CoverLetterContent) -> bytes:
        stream = BytesIO()
        document = SimpleDocTemplate(
            stream,
            pagesize=letter,
            rightMargin=0.9 * inch,
            leftMargin=0.9 * inch,
            topMargin=0.85 * inch,
            bottomMargin=0.85 * inch,
            title="Cover Letter",
            author="",
        )
        styles = getSampleStyleSheet()
        body = ParagraphStyle(
            "CoverLetterBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#10231F"),
            spaceAfter=10,
        )
        name = ParagraphStyle(
            "CoverLetterName",
            parent=body,
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=18,
            spaceAfter=2,
        )
        contact = ParagraphStyle(
            "CoverLetterContact",
            parent=body,
            fontSize=9,
            leading=11,
            textColor=colors.HexColor("#4D5E59"),
            spaceAfter=14,
        )
        role = ParagraphStyle(
            "CoverLetterRole",
            parent=body,
            fontName="Helvetica-Bold",
            spaceAfter=14,
        )

        story = [Paragraph(escape(content.candidate_name), name)]
        if content.candidate_email:
            story.append(Paragraph(escape(content.candidate_email), contact))
        story.append(
            Paragraph(
                escape(f"{content.job_title} at {content.company_name}"),
                role,
            )
        )
        for paragraph_text in (
            content.salutation,
            content.opening,
            content.evidence_paragraph,
            content.motivation_paragraph,
            content.closing_paragraph,
            content.sign_off,
            content.candidate_name,
        ):
            story.extend([Paragraph(escape(paragraph_text), body), Spacer(1, 1)])
        document.build(story)
        return stream.getvalue()
